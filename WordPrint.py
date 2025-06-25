from docx import Document
import ExcelExtrator as EE
from docx.shared import Pt

QOUALCELL = 4
STOPCOL = 3
COLS = 11
FIRSTCOL = 0
SECONDTABLE = 1
FONTSIZE = 12
MASTAB = 2

doc = Document('tofestov.docx')
filename = "tofes.xlsx"


def checkQoual(excelFile, row_index, col_index):
    try:
        cell_value = excelFile.iloc[row_index, col_index]
        return str(cell_value).strip() == 'כן'
    except Exception as e:
        print(f"שגיאה בקריאת תא: {e}")
        return False


#מדפיס את תא האקסל לתא נבחר
def wordPrint(doc, df, row_index, col_index , increase):
    table = doc.tables[SECONDTABLE]
    cell = table.cell(row_index, col_index )
    paragraph = cell.paragraphs[FIRSTCOL]
    run = paragraph.add_run(EE.readExcel(df, row_index, col_index + increase))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(FONTSIZE)
    print("עודכן בהצלחה!")


def excelEctract(doc, df):
    for i in range(len(df)):
        wordPrint(doc, df, i,FIRSTCOL,MASTAB )
        if not checkQoual(df, i, QOUALCELL):
            for j in range(1,COLS):
                wordPrint(doc, df, i, j,STOPCOL)

